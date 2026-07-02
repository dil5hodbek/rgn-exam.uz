import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from app.services.exercise_registry import default_kind
from app.services.question_classifier import (
    clean_text,
    extract_alternative_pair,
    extract_labelled_options,
    infer_type,
)

SECTION_NAMES = {
    "listening", "grammar", "vocabulary", "function", "reading", "writing",
    "use of english", "speaking", "pronunciation",
}
TASK_START = re.compile(
    r"^\s*(\d{1,2})\s*(?:[|.)-]\s*)?"
    r"(recording|listen|read|complete|choose|write|match|correct|underline|put|tick|look|fill|rewrite|"
    r"answer|circle|number|decide|find|select|replace|order)\b",
    re.I,
)
ITEM_START = re.compile(r"^\s*(\d{1,2})\s*(?:[|.)-]\s*)?(.+)", re.S)
TEST_HEADING = re.compile(
    r"^\s*(MID[\s-]*COURSE|END(?:[\s-]*OF)?[\s-]*COURSE)\s+TEST(?:\s+V)?\s*(\d+)\s*$",
    re.I,
)
EXERCISE_HEADING = re.compile(r"^\s*Exercise\s+(\d+)\s*$", re.I)


def slug_exam(heading: str) -> str:
    return "mid-course" if heading.lower().startswith("mid") else "end-course"


@dataclass
class ParsedQuestion:
    number: int
    prompt: str
    options: list[str] = field(default_factory=list)
    correct_answer: Any | None = None
    accepted_answers: list[Any] = field(default_factory=list)
    is_example: bool = False


@dataclass
class ParsedTask:
    exercise_number: int
    title: str
    instructions: str
    question_type: str
    section: str
    recording_number: int | None = None
    passage: str | None = None
    questions: list[ParsedQuestion] = field(default_factory=list)
    interaction: dict[str, Any] = field(default_factory=dict)
    confidence: float = 1.0
    warnings: list[str] = field(default_factory=list)


@dataclass
class ParsedTest:
    level: str
    exam_slug: str
    variant: int
    title: str
    source_path: Path
    tasks: list[ParsedTask]
    warnings: list[str] = field(default_factory=list)


class AnswerKey:
    def __init__(self, path: Path | None = None):
        self.answers: dict[tuple[str, int, int], dict[int, str]] = {}
        self.raw_blocks: dict[tuple[str, int, int], str] = {}
        if path is not None:
            self._parse(path)

    @classmethod
    def empty(cls) -> "AnswerKey":
        """An answer key with no entries, for archives where no key file was found."""
        return cls(path=None)

    @staticmethod
    def _answer_pairs(lines: list[str]) -> dict[int, str]:
        result: dict[int, str] = {}
        for raw in lines:
            line = clean_text(raw).replace("**", "")
            if not line or line.casefold().startswith(("guide for", "use the following", "sample answer", "varies")):
                continue
            matches = list(re.finditer(
                r"(?:^|\s)(\d{1,2})\s*(?:[.)-]\s*)?(.+?)(?=(?:\s+\d{1,2}\s*(?:[.)-]\s*)?)|$)",
                line,
            ))
            for match in matches:
                number = int(match.group(1))
                answer = clean_text(match.group(2)).strip(" |;,")
                if answer and len(answer) < 180:
                    result[number] = answer
        return result

    def _parse(self, path: Path) -> None:
        paragraphs = [
            clean_text(line)
            for paragraph in Document(path).paragraphs
            for line in paragraph.text.splitlines()
            if clean_text(line)
        ]
        exam_slug: str | None = None
        variant: int | None = None
        exercise: int | None = None
        block: list[str] = []

        def commit() -> None:
            if exam_slug and variant and exercise and block:
                key = (exam_slug, variant, exercise)
                self.answers[key] = self._answer_pairs(block)
                self.raw_blocks[key] = "\n".join(block)

        for text in paragraphs:
            test_match = TEST_HEADING.match(text)
            if test_match:
                commit()
                block = []
                exam_slug, variant, exercise = slug_exam(test_match.group(1)), int(test_match.group(2)), None
                continue
            exercise_match = EXERCISE_HEADING.match(text)
            if exercise_match and exam_slug and variant:
                commit()
                block = []
                exercise = int(exercise_match.group(1))
                continue
            if text.casefold() in SECTION_NAMES:
                continue
            if exercise:
                block.append(text)
        commit()

    def for_exercise(self, exam_slug: str, variant: int, exercise: int) -> dict[int, str]:
        return self.answers.get((exam_slug, variant, exercise), {})


def extract_lettered_paragraphs(lines: list[str]) -> list[dict[str, str]]:
    """Find a continuous A-F article block without confusing a/b/c answer options."""
    candidates: list[dict[str, str]] = []
    current_value: str | None = None
    current_lines: list[str] = []
    expected = "A"
    for line in lines:
        marker = line.strip().upper()
        if marker == expected:
            if current_value:
                candidates.append({"value": current_value.casefold(), "label": " ".join(current_lines).strip()})
            current_value, current_lines = marker, []
            expected = chr(ord(expected) + 1)
            continue
        if current_value:
            if len(marker) == 1 and marker.isalpha() and marker != expected:
                break
            current_lines.append(line)
    if current_value:
        candidates.append({"value": current_value.casefold(), "label": " ".join(current_lines).strip()})
    return [item for item in candidates if item["label"]] if len(candidates) >= 3 else []


def inline_alternative_questions(
    raw_body: str,
    answers: dict[int, str],
) -> tuple[list[ParsedQuestion], dict[str, Any]] | None:
    markers = list(re.finditer(r"(?<!\d)(\d{1,2})(?=[^\d\s])", raw_body))
    if len(markers) < 2:
        return None
    questions: list[ParsedQuestion] = []
    items: list[dict[str, Any]] = []
    for index, marker in enumerate(markers):
        number = int(marker.group(1))
        end = markers[index + 1].start() if index + 1 < len(markers) else len(raw_body)
        segment = raw_body[marker.end():end]
        slash = segment.find("/")
        if slash < 1:
            continue
        left = segment[:slash].strip(" |")
        if len(left.split()) > 5:
            continue
        after = segment[slash + 1:].lstrip()
        right_match = re.match(r"([\w'’-]+(?:\s+[\w'’-]+){0,4})", after, re.UNICODE)
        if not right_match:
            continue
        right_words = right_match.group(1).split()
        correct = answers.get(number)
        right = ""
        if correct:
            normalized_correct = clean_text(str(correct)).casefold()
            for length in range(1, len(right_words) + 1):
                candidate = " ".join(right_words[:length]).strip(" |,.;:?!")
                if candidate.casefold() == normalized_correct:
                    right = candidate
                    break
        if not right:
            left_words = left.split()
            length = min(len(left_words), len(right_words))
            if left.casefold().startswith("to ") and right_words:
                length = 1
            right = " ".join(right_words[:length]).strip(" |,.;:?!")
        consumed = right_match.start(1) + len(right)
        after_text = after[consumed:]
        prompt = f"{left} / {right}{after_text}".strip()
        is_example = correct is None
        selected = left if is_example else clean_text(str(correct))
        questions.append(ParsedQuestion(
            number=number,
            prompt=prompt,
            options=[left, right],
            correct_answer=None if is_example else selected,
            is_example=is_example,
        ))
        items.append({
            "number": number,
            "before": "",
            "options": [left, right],
            "after": after_text,
            "example_value": left if is_example else None,
        })
    if len(questions) < 2:
        return None
    return questions, {"kind": "inline_alternatives", "items": items}


class RoadmapDocumentParser:
    def __init__(self, answer_key: AnswerKey):
        self.answer_key = answer_key

    def parse(
        self, path: Path, level: str | None, exam_slug: str | None, variant: int | None,
    ) -> ParsedTest:
        document = Document(path)
        paragraphs = [clean_text(p.text) for p in document.paragraphs if clean_text(p.text)]
        for table in document.tables:
            for row in table.rows:
                value = " | ".join(clean_text(cell.text) for cell in row.cells if clean_text(cell.text))
                if value:
                    paragraphs.append(value)
        lettered_paragraphs = extract_lettered_paragraphs(paragraphs)

        section = "General"
        chunks: list[tuple[str, int, str, list[str]]] = []
        current_number: int | None = None
        current_title = ""
        current_body: list[str] = []

        def commit() -> None:
            if current_number is not None:
                chunks.append((section, current_number, current_title, list(current_body)))

        for text in paragraphs:
            if text.casefold() in SECTION_NAMES:
                commit()
                current_number, current_title, current_body = None, "", []
                section = text.title()
                continue
            task_match = TASK_START.match(text)
            if task_match:
                commit()
                current_number = int(task_match.group(1))
                current_title = text
                current_body = []
                continue
            if current_number is not None and not re.fullmatch(r"/\s*\d+", text):
                current_body.append(text)
        commit()

        tasks: list[ParsedTask] = []
        for exercise_number, (section_name, printed_number, title, body) in enumerate(chunks, start=1):
            # Printed exercise numbers reset visually in a few converted documents.
            actual_exercise = printed_number if printed_number >= exercise_number else exercise_number
            question_type = infer_type(title, section_name)
            # Answer keys follow the printed number, even when a converted DOCX
            # contains a duplicate number and the student-facing order is adjusted.
            answers = self.answer_key.for_exercise(exam_slug, variant, printed_number)
            recording_match = re.search(r"Recording\s*(\d+)", title, re.I)
            task = ParsedTask(
                exercise_number=actual_exercise,
                title=f"Exercise {actual_exercise}",
                instructions=re.sub(r"^\s*\d+\s*(?:[|.)-]\s*)?", "", title).strip(),
                question_type=question_type,
                section=section_name,
                recording_number=int(recording_match.group(1)) if recording_match else None,
            )
            task.questions, task.passage, task.interaction = self._questions(
                body, question_type, answers, section_name, task.instructions,
            )
            if question_type == "matching_headings":
                task.interaction = {
                    "kind": "matching_headings",
                    "options": lettered_paragraphs or [
                        {"value": letter, "label": f"Paragraph {letter.upper()}"}
                        for letter in "abcdef"
                    ],
                    "reuse_options": False,
                }
                if lettered_paragraphs:
                    task.passage = "\n\n".join(
                        f"{item['value'].upper()}\n{item['label']}" for item in lettered_paragraphs
                    )
                else:
                    task.confidence = min(task.confidence, 0.7)
                    task.warnings.append("Matching headings paragraphs could not be extracted.")
                for question in task.questions:
                    if question.correct_answer:
                        match = re.match(r"([a-fA-F])", str(question.correct_answer))
                        question.correct_answer = match.group(1).casefold() if match else question.correct_answer
            if not task.questions:
                task.questions = [ParsedQuestion(
                    number=1,
                    prompt="\n".join(body) or task.instructions,
                    correct_answer=None,
                )]
                task.confidence = 0.35
                if task.question_type == "multiple_choice" and re.search(
                    r"[\w'’-]+\s*/\s*[\w'’-]+", task.questions[0].prompt,
                ):
                    task.interaction = {"kind": "inline_alternatives"}
                elif task.question_type == "writing":
                    task.interaction = {"kind": "long_text", "manual_review": True}
                    task.confidence = 0.9
            if task.question_type == "rich_text_question":
                task.interaction = {"kind": "long_text", "manual_review": True}
                task.confidence = min(task.confidence, 0.65)
                task.warnings.append("No safe automatic grading strategy was detected.")
            tasks.append(task)
        tasks = self._merge_writing_workflows(tasks)
        display_level = level or "Unspecified"
        display_exam = "Mid-course" if exam_slug == "mid-course" else "End-course"
        display_variant = variant if variant is not None else 1
        return ParsedTest(
            level=display_level,
            exam_slug=exam_slug or "general",
            variant=display_variant,
            title=f"{display_level} {display_exam} Test {display_variant}",
            source_path=path,
            tasks=tasks,
        )

    @staticmethod
    def _merge_writing_workflows(tasks: list[ParsedTask]) -> list[ParsedTask]:
        result: list[ParsedTask] = []
        index = 0
        while index < len(tasks):
            task = tasks[index]
            if (
                task.section.casefold() == "writing"
                and re.search(r"\bwrite an essay\b", task.instructions, re.I)
            ):
                group = [task]
                cursor = index + 1
                while cursor < len(tasks) and tasks[cursor].section.casefold() == "writing" and re.search(
                    r"\b(make notes|organise your ideas|write \d+\s*[–-]\s*\d+ words)\b",
                    tasks[cursor].instructions,
                    re.I,
                ):
                    group.append(tasks[cursor])
                    cursor += 1
                if len(group) > 1:
                    instructions = "\n".join(
                        f"{position + 1}. {item.instructions}" for position, item in enumerate(group)
                    )
                    range_match = re.search(r"(\d+)\s*[–-]\s*(\d+)\s*words", instructions, re.I)
                    task.instructions = instructions
                    task.question_type = "writing"
                    task.interaction = {
                        "kind": "long_text",
                        "manual_review": True,
                        "min_words": int(range_match.group(1)) if range_match else None,
                        "max_words": int(range_match.group(2)) if range_match else None,
                        "workflow_steps": [item.instructions for item in group],
                        "merged_exercises": [item.exercise_number for item in group],
                    }
                    task.questions = [ParsedQuestion(
                        number=1,
                        prompt=group[0].questions[0].prompt if group[0].questions else group[0].instructions,
                        correct_answer=None,
                    )]
                    task.confidence = max(0.9, min(item.confidence for item in group))
                    result.append(task)
                    index = cursor
                    continue
            result.append(task)
            index += 1
        return result

    def _questions(
        self,
        body: list[str],
        question_type: str,
        answers: dict[int, str],
        section: str,
        instructions: str,
    ) -> tuple[list[ParsedQuestion], str | None, dict[str, Any]]:
        blocks: list[tuple[int, list[str]]] = []
        intro: list[str] = []
        number: int | None = None
        current: list[str] = []
        for text in body:
            match = ITEM_START.match(text)
            if match and int(match.group(1)) <= 30:
                if number is not None:
                    blocks.append((number, current))
                number = int(match.group(1))
                current = [match.group(2)]
            elif number is None:
                intro.append(text)
            else:
                current.append(text)
        if number is not None:
            blocks.append((number, current))

        raw_body = "\n".join(body)
        embedded_alternatives = inline_alternative_questions(raw_body, answers)
        if embedded_alternatives:
            questions, interaction = embedded_alternatives
            return questions, None, interaction

        embedded_gap_pattern = re.compile(
            r"\b(\d{1,2})\s*(_{2,}(?:[^\W\d_]+_{2,})?)",
            re.UNICODE,
        )
        embedded_gaps = list(embedded_gap_pattern.finditer(raw_body))

        def sentence_context(gap: re.Match[str], item_number: int) -> str:
            sentence_start = max(raw_body.rfind(".", 0, gap.start()) + 1, raw_body.rfind("\n", 0, gap.start()) + 1)
            sentence_end = raw_body.find(".", gap.end())
            sentence_end = len(raw_body) if sentence_end < 0 else sentence_end + 1
            context = raw_body[sentence_start:sentence_end].strip()
            return embedded_gap_pattern.sub(
                lambda match: "_____" if int(match.group(1)) == item_number else match.group(0),
                context,
            )

        # Each gap has its own a/b/c options listed on its own line right
        # after the passage, e.g. "1 a issue b present c habit" — a cloze
        # passage where every blank is itself a small multiple-choice.
        gap_choice_blocks: dict[int, list[dict[str, str]]] = {}
        for item_number, lines in blocks:
            labelled, options_start = extract_labelled_options("\n".join(lines))
            # A genuine per-gap choice line has short option labels; a huge
            # label means the block absorbed trailing content past the last
            # option (e.g. an unrelated passage with no exercise boundary
            # between it and this block) — discard rather than keep noise.
            if len(labelled) >= 2 and (options_start or 0) <= 2 and all(len(option["label"]) <= 40 for option in labelled):
                gap_choice_blocks[item_number] = labelled
        if len(embedded_gaps) >= 3 and len(gap_choice_blocks) >= max(2, len(embedded_gaps) - 2):
            template = embedded_gap_pattern.sub(lambda match: f"{{{{{match.group(1)}}}}}", raw_body)
            questions: list[ParsedQuestion] = []
            example_values: dict[str, str] = {}
            for gap in embedded_gaps:
                item_number = int(gap.group(1))
                filled_value = gap.group(2).strip("_")
                is_example = bool(filled_value)
                if is_example:
                    example_values[str(item_number)] = filled_value
                options = gap_choice_blocks.get(item_number, [])
                option_labels = [option["label"] for option in options]
                answer_raw = answers.get(item_number)
                answer: str | None = None
                if answer_raw and not is_example:
                    letter = re.match(r"^([a-hA-H])(?:\b|[).])", str(answer_raw).strip())
                    if letter and option_labels:
                        option_index = ord(letter.group(1).casefold()) - ord("a")
                        if 0 <= option_index < len(option_labels):
                            answer = option_labels[option_index]
                    if answer is None:
                        answer = re.sub(r"^[a-hA-H][.)]\s*", "", str(answer_raw)).strip()
                questions.append(ParsedQuestion(
                    number=item_number,
                    prompt=clean_text(sentence_context(gap, item_number)),
                    options=option_labels,
                    correct_answer=None if is_example else answer,
                    is_example=is_example,
                ))
            questions.sort(key=lambda item: item.number)
            return questions, None, {
                "kind": "cloze_passage",
                "template": template,
                "example_values": example_values,
                "per_question_options": True,
            }

        if len(blocks) <= 1 and len(embedded_gaps) >= 3:
            template = embedded_gap_pattern.sub(lambda match: f"{{{{{match.group(1)}}}}}", raw_body)
            questions: list[ParsedQuestion] = []
            example_values: dict[str, str] = {}
            for gap in embedded_gaps:
                item_number = int(gap.group(1))
                filled_value = gap.group(2).strip("_")
                sentence_start = max(raw_body.rfind(".", 0, gap.start()) + 1, raw_body.rfind("\n", 0, gap.start()) + 1)
                sentence_end = raw_body.find(".", gap.end())
                if sentence_end < 0:
                    sentence_end = len(raw_body)
                else:
                    sentence_end += 1
                context = raw_body[sentence_start:sentence_end].strip()
                context = embedded_gap_pattern.sub(
                    lambda match: "_____" if int(match.group(1)) == item_number else match.group(0),
                    context,
                )
                is_example = bool(filled_value)
                if is_example:
                    example_values[str(item_number)] = filled_value
                answer = answers.get(item_number)
                alternatives = [
                    part.strip() for part in re.split(r"\s*/\s*", str(answer)) if part.strip()
                ] if answer and "/" in str(answer) else []
                questions.append(ParsedQuestion(
                    number=item_number,
                    prompt=clean_text(context),
                    correct_answer=None if is_example else (alternatives[0] if alternatives else answer),
                    accepted_answers=alternatives[1:] if alternatives else [],
                    is_example=is_example,
                ))
            return questions, None, {
                "kind": "cloze_passage",
                "template": template,
                "example_values": example_values,
            }

        questions: list[ParsedQuestion] = []
        for index, (item_number, lines) in enumerate(blocks):
            combined = "\n".join(lines)
            options: list[str] = []
            if question_type in {"multiple_choice", "matching_pairs", "matching_headings", "sentence_ordering"}:
                labelled, options_start = extract_labelled_options(combined)
                options = [option["label"] for option in labelled]
                if options_start is not None:
                    combined = combined[:options_start].strip()
            answer = answers.get(item_number)
            if answer:
                answer = re.sub(r"^[a-hA-H][.)]\s*", "", answer).strip()
                if question_type == "true_false":
                    answer = {"t": "True", "f": "False"}.get(answer.casefold(), answer)
                elif question_type == "true_false_not_given":
                    answer = {
                        "t": "True", "f": "False", "n": "Not Given", "ni": "Not Given",
                    }.get(answer.casefold(), answer)
                elif question_type == "multiple_choice":
                    letter = re.match(r"^([a-hA-H])(?:\b|$)", answer)
                    if letter and options:
                        option_index = ord(letter.group(1).casefold()) - ord("a")
                        if 0 <= option_index < len(options):
                            answer = options[option_index]
            questions.append(ParsedQuestion(
                number=item_number,
                prompt=clean_text(combined),
                options=options,
                correct_answer=answer,
                accepted_answers=[part.strip() for part in re.split(r"\s*/\s*", answer) if part.strip()] if answer and "/" in answer else [],
                is_example=(
                    question_type not in {"writing", "rich_text_question"}
                    and index == 0
                    and (
                        item_number not in answers
                        or bool(re.search(r"_{1,}\s*[\w'’-]+\s*_{1,}|\|\s*[A-HLETF]\s*$", combined, re.I))
                    )
                ),
            ))
        passage = "\n\n".join(intro) if section.casefold() == "reading" and intro else None
        interaction: dict[str, Any] = {}
        instruction_text = instructions.casefold()
        intro_text = " ".join(intro)
        labelled_source = " ".join(dict.fromkeys([
            *intro,
            *(text for text in body if re.match(r"^\s*[a-hA-H](?:[.)-]\s*|\s+)", text)),
        ]))
        labelled_options, _ = extract_labelled_options(labelled_source)

        if question_type == "gap_fill" and "box" in instruction_text and intro_text:
            answer_values = [
                str(answer).strip()
                for answer in answers.values()
                if answer and not re.fullmatch(r"[a-hA-H]", str(answer).strip())
            ]
            if "|" in intro_text:
                words = [
                    option.strip(" ,;|")
                    for option in re.split(r"\s*\|\s*", intro_text)
                    if option.strip(" ,;|")
                ]
            else:
                answer_pattern = (
                    re.compile("(" + "|".join(re.escape(answer) for answer in sorted(set(answer_values), key=len, reverse=True)) + ")", re.I)
                    if answer_values else None
                )
                pieces = answer_pattern.split(intro_text) if answer_pattern else [intro_text]
                words = []
                for piece in pieces:
                    if not piece.strip():
                        continue
                    matched_answer = next((answer for answer in answer_values if piece.casefold() == answer.casefold()), None)
                    if matched_answer:
                        words.append(matched_answer)
                    else:
                        words.extend(word.strip(" ,;|") for word in piece.split() if word.strip(" ,;|"))
            interaction = {
                "kind": "word_bank",
                "options": list(dict.fromkeys(words)),
                "reuse_options": bool(re.search(r"\b(use|using)\s+(?:one|a)\s+word\s+twice\b", instruction_text)),
            }
        elif question_type in {"matching_pairs", "gap_fill"} and labelled_options:
            interaction = {
                "kind": "matching",
                "options": labelled_options,
                "reuse_options": bool(re.search(
                    r"\b(speakers?|may belong|more than once|can be used twice|two opinions)\b",
                    instruction_text,
                )),
            }
        elif question_type == "matching_pairs" and re.search(
            r"\([A-Za-z]\)\s*[A-Za-z][\w\s]*?\s+and\s+\([A-Za-z]\)",
            instructions,
        ):
            # Generic "find X of (A) label-one and (B) label-two" two-category
            # matching, e.g. "...consequences of (L) little sleep and (E) enough sleep".
            category_match = re.search(
                r"\(([A-Za-z])\)\s*([A-Za-z][\w\s]*?)\s+and\s+\(([A-Za-z])\)\s*([A-Za-z][\w\s]*?)(?:[.,;:]|$)",
                instructions,
            )
            if category_match:
                first_value, first_label, second_value, second_label = category_match.groups()
                options = [
                    {"value": first_value.casefold(), "label": first_label.strip().capitalize()},
                    {"value": second_value.casefold(), "label": second_label.strip().capitalize()},
                ]
            else:
                options = [
                    {"value": "l", "label": "Little sleep"},
                    {"value": "e", "label": "Enough sleep"},
                ]
            interaction = {"kind": "matching", "options": options, "reuse_options": True}
            for question in questions:
                if question.correct_answer:
                    question.correct_answer = str(question.correct_answer).strip().casefold()
        elif question_type == "sentence_ordering":
            interaction = {"kind": "ordering"}
        elif question_type == "multiple_choice" and (
            "alternative" in instruction_text
            or any(re.search(r"[\w'’-]+\s*/\s*[\w'’-]+", question.prompt) for question in questions)
        ):
            items = []
            for question in questions:
                pair = extract_alternative_pair(question.prompt, question.correct_answer)
                if pair:
                    question.options = [pair["left"], pair["right"]]
                    items.append({
                        "number": question.number,
                        "before": pair["before"],
                        "options": question.options,
                        "after": pair["after"],
                    })
            interaction = {"kind": "inline_alternatives", "items": items}
        elif question_type in {"true_false", "true_false_not_given"}:
            interaction = {"kind": "binary_choice"}
        elif question_type == "multiple_choice":
            interaction = {"kind": "multiple_choice"}
        elif question_type == "gap_fill":
            interaction = {"kind": "guided_input"}
        elif question_type == "error_correction":
            interaction = {"kind": "correction"}
        elif question_type == "short_answer":
            interaction = {"kind": "short_answer"}
        elif question_type == "writing":
            range_match = re.search(r"(\d+)\s*[–-]\s*(\d+)\s*words", instructions, re.I)
            interaction = {
                "kind": "long_text",
                "manual_review": True,
                "min_words": int(range_match.group(1)) if range_match else None,
                "max_words": int(range_match.group(2)) if range_match else None,
            }
        else:
            interaction = {"kind": default_kind(question_type), "manual_review": True}

        return questions, passage, interaction


def discover_tests(root: Path) -> list[tuple[Path, str, str, int]]:
    result: list[tuple[Path, str, str, int]] = []
    for path in root.rglob("*.docx"):
        lower = path.name.casefold()
        if path.name.startswith("~") or "answer" in lower:
            continue
        joined = "/".join(path.parts).casefold()
        level = next((name for name in ("Pre-Intermediate", "Intermediate", "Elementary", "Beginner") if name.casefold() in joined), None)
        exam_slug = "mid-course" if "/mid/" in joined.replace("\\", "/") else "end-course" if "/end/" in joined.replace("\\", "/") else None
        variant_match = re.search(r"\bV\s*(\d+)\b", path.stem, re.I)
        if level and exam_slug and variant_match:
            result.append((path, level, exam_slug, int(variant_match.group(1))))
    return sorted(result, key=lambda item: (item[1], item[2], item[3]))
