"""Best-effort .docx → exercise parser.

Extracts *question text only* (prompts and, for multiple choice, the options),
guesses the question type, and returns a structure the admin exercise builder
can pre-fill. It deliberately does NOT guess correct answers or media — the
admin sets those by hand.

Two strategies, in order:

1. **AI** (:func:`_ai_parse`) — if ``OPENROUTER_API_KEY`` (preferred) or
   ``ANTHROPIC_API_KEY`` is configured, the raw document text is sent to an
   LLM, which classifies the question type and splits the questions/options
   far more reliably than regexes across the many real coursebook layouts.
   Only *question text* is asked for; answers/media stay manual.
2. **Heuristics** (:func:`_parse_repeat` / :func:`_parse_matching`) — the
   original regex parser, used automatically when there is no API key, the SDK
   is missing, or the AI call fails/returns nothing. So import always works.

The heuristics handle the common coursebook layout where:
* the exercise header starts with a number + rubric ("1 Recording 1 Listen …");
* questions start with a number, with or without punctuation ("1 Fiona is …");
* options are lettered a/b/c, either one per line ("a The Lion …") or grouped
  on one line ("a buy a book  b the shops  c Matt's house").
"""

import asyncio
import html as html_escape
import json
import logging
import re
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)

# A question line: number, optional . ) - : delimiter, then the text.
QNUM_RE = re.compile(r"^\(?(\d{1,3})\s*[\.\)\-:]?\s+(.+)$")
# Option letters are lowercase a–h. (Uppercase A:/B: are dialogue speakers, not
# options — keeping options lowercase avoids mistaking a dialogue turn for one.)
SINGLE_OPT_DELIM_RE = re.compile(r"^\(?([a-h])\s*[\.\)\-:]\s+(.+)$")
LONE_OPT_RE = re.compile(r"^\(?([a-h])\)?\s+(.+)$")
LETTER_TOKEN_RE = re.compile(r"(?:^|\s)([a-h])(?=\s)")

INSTRUCTION_PHRASES = (
    "recording", "listen to", "listen and", "choose the", "choose from",
    "complete the", "complete each", "match the", "decide if", "decide whether",
    "true or false", "not given", "fill in", "fill the", "underline the",
    "circle the", "read the text", "read the passage", "read the article",
    "put the", "write the", "rewrite the", "look at the", "answer the questions",
    "tick the", "correct the",
)


def _extract_lines(data: bytes) -> list[str]:
    from docx import Document  # lazy import so the module loads without the dep
    document = Document(BytesIO(data))
    lines: list[str] = []
    seen: set[str] = set()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen:
                    seen.add(text)
                    lines.append(text)
    return lines


def _looks_instruction(text: str) -> bool:
    low = text.lower()
    return any(phrase in low for phrase in INSTRUCTION_PHRASES)


def _split_inline_options(text: str) -> tuple[str, list[str]]:
    """Turn 'a buy a book  b the shops  c Matt's house' into
    ('', ['buy a book', 'the shops', "Matt's house"]).

    The option letters run in order a, b, c, … so we lock onto that sequence —
    that way a stray 'a' inside an option's text (e.g. 'buy a book') is ignored.
    Returns (prefix_before_first_option, options)."""
    marks = [(m.start(1), m.group(1).lower()) for m in LETTER_TOKEN_RE.finditer(text)]
    letters = "abcdefghij"
    seq: list[int] = []
    want = 0
    for pos, ch in marks:
        if want < len(letters) and ch == letters[want]:
            seq.append(pos)
            want += 1
    if len(seq) < 2:
        return text, []
    prefix = text[: seq[0]].strip(" .,:;-")
    options: list[str] = []
    for i, pos in enumerate(seq):
        start = pos + 1
        end = seq[i + 1] if i + 1 < len(seq) else len(text)
        options.append(text[start:end].strip(" .,;"))
    return prefix, [opt for opt in options if opt]


def _guess_type(instructions: str, questions: list[dict[str, Any]]) -> str:
    instr = instructions.lower()
    has_options = any(q["options"] for q in questions)
    has_gap = any(("___" in q["prompt"] or "…" in q["prompt"] or "....." in q["prompt"]) for q in questions)
    # "word1 / word2" choices inside the sentences → inline alternatives.
    has_slash = any(re.search(r"\w\s*/\s*\w", q["prompt"]) for q in questions)

    if "alternative" in instr or (has_slash and not has_options):
        return "inline_alternatives"
    if "correct the mistake" in instr or "correct the error" in instr \
            or "find the mistake" in instr or "rewrite the correct" in instr:
        return "error_correction"
    if "true" in instr and "false" in instr:
        return "true_false_not_given" if "not given" in instr else "true_false"
    if "match" in instr:
        return "matching_pairs"
    if has_options:
        return "multiple_choice"
    if has_gap or "gap" in instr or "fill" in instr:
        return "gap_fill"
    if "short" in instr or "answer the" in instr:
        return "short_answer"
    return "gap_fill"


LETTERED_LINE_RE = re.compile(r"^\(?([a-h])\)?\s*[\.\)\-:]?\s+(.+)$")


def _parse_matching(lines: list[str], instructions: str) -> dict[str, Any]:
    """'Match 1–8 with a–h': numbered items are the left prompts, lettered items
    are the right options (replies). The admin sets the correct pairs."""
    left: list[str] = []
    right: list[str] = []
    collected = False
    for line in lines:
        if re.fullmatch(r"[\s_\.\-–—…]+", line):
            continue
        # Skip the leading rubric/header only — once real items start, a question
        # or reply that happens to contain a phrase like "listen to" is content.
        if not collected and _looks_instruction(line):
            continue
        qn = QNUM_RE.match(line)
        if qn:
            left.append(qn.group(2).strip())
            collected = True
            continue
        lettered = LETTERED_LINE_RE.match(line)
        if lettered:
            right.append(lettered.group(2).strip())
            collected = True
            continue
        # Continuation of the last collected item (or an ignored word box before any).
        if right:
            right[-1] = f"{right[-1]} {line}".strip()
        elif left:
            left[-1] = f"{left[-1]} {line}".strip()
    left = [re.sub(r"_{2,}", "___", x) for x in left if x]
    right = [x for x in right if x]
    return {"template_key": "matching_pairs", "instructions": instructions, "left": left, "right": right, "questions": []}


# ---------------------------------------------------------------------------
# AI strategy
# ---------------------------------------------------------------------------

# The exercise types the AI is allowed to choose. Limited to the ones the
# import flow can pre-fill cleanly: per-question "repeat" types + matching_pairs
# (left/right columns). Composite types with bespoke editors (ordering, cloze,
# heading-matching) are excluded — the model picks the closest supported type,
# or the admin selects it by hand.
_AI_ALLOWED = {
    "multiple_choice", "multi_select", "true_false", "true_false_not_given",
    "gap_fill", "dropdown_gap_fill", "short_answer", "error_correction",
    "inline_alternatives", "writing", "speaking_prompt", "matching_pairs",
    "gap_match",
}
_MATCHING_KEYS = {"matching_pairs"}

_AI_SYSTEM_PROMPT = """You extract exercises from an English coursebook page and return STRICT JSON.

You are given the raw text of ONE exercise from an English test/coursebook. Your job:
1. Identify the single best question TYPE (template_key) from this exact list:
   multiple_choice, multi_select, true_false, true_false_not_given, gap_fill,
   dropdown_gap_fill, short_answer, error_correction, inline_alternatives,
   writing, speaking_prompt, matching_pairs
2. Split the exercise into individual questions and pull out ONLY the question
   text (and, for choice questions, the answer options).

Rules:
- Extract ONLY question/prompt text and answer options. NEVER invent, guess, or
  fill in the correct answer. NEVER include audio/media references, page
  numbers, "Recording 1", track numbers, or the answer key.
- Strip leading numbering ("1", "2)", "a.") from prompts and options.
- Options are the choices only, without their a/b/c letters.
- Keep any blank/gap in a prompt as exactly three underscores: ___
- "instructions" is the exercise rubric ("Listen and choose the correct
  answer", "Complete the sentences", "Match ...", etc.), cleaned of leading
  numbers. Empty string if there is no rubric.
- For inline_alternatives, put the two choices inside the prompt as
  "word1 / word2" and leave options empty.
- For "odd one out"/"choose the different word" groups, use multiple_choice:
  prompt is the instruction and options are the words in that group.

Output shape:
- For matching_pairs:
  {"template_key": "matching_pairs", "instructions": "...",
   "left": ["left item 1", "left item 2", ...],
   "right": ["right option 1", "right option 2", ...],
   "questions": []}
- For every other type:
  {"template_key": "gap_fill", "instructions": "...",
   "questions": [{"prompt": "...", "options": ["...", "..."]}]}
  (options is [] when the type has no choices)

Return ONLY the JSON object. No markdown, no code fences, no commentary."""


def _extract_json(raw: str) -> dict[str, Any] | None:
    """Pull the JSON object out of the model reply, tolerating code fences or
    stray prose around it."""
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end > start:
        try:
            return json.loads(text[start : end + 1])
        except json.JSONDecodeError:
            return None
    return None


def _clean_str(value: Any) -> str:
    s = str(value).strip() if value is not None else ""
    return re.sub(r"_{2,}", "___", s)


def _normalise_ai(data: dict[str, Any]) -> dict[str, Any] | None:
    """Validate and coerce the model output into the exact shape the builder
    expects. Returns None if it is unusable so we fall back to heuristics."""
    if not isinstance(data, dict):
        return None
    template_key = str(data.get("template_key") or "").strip()
    if template_key not in _AI_ALLOWED:
        return None
    instructions = _clean_str(data.get("instructions"))

    if template_key in _MATCHING_KEYS:
        left = [_clean_str(x) for x in (data.get("left") or []) if _clean_str(x)]
        right = [_clean_str(x) for x in (data.get("right") or []) if _clean_str(x)]
        if len(left) < 2 or len(right) < 2:
            return None
        return {
            "template_key": template_key,
            "instructions": instructions,
            "left": left,
            "right": right,
            "questions": [],
        }

    questions: list[dict[str, Any]] = []
    for item in data.get("questions") or []:
        if not isinstance(item, dict):
            continue
        prompt = _clean_str(item.get("prompt"))
        if not prompt:
            continue
        options = [_clean_str(o) for o in (item.get("options") or []) if _clean_str(o)]
        questions.append({"prompt": prompt, "options": options})
    if not questions:
        return None
    return {
        "template_key": template_key,
        "instructions": instructions,
        "questions": questions,
    }


def _openrouter_complete(api_key: str, model: str, text: str, max_tokens: int = 4096) -> str | None:
    """Call OpenRouter's OpenAI-compatible chat completions API. Returns the
    reply text, or None on any failure."""
    import httpx

    try:
        response = httpx.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": model,
                "max_tokens": max_tokens,
                "messages": [
                    {"role": "system", "content": _AI_SYSTEM_PROMPT},
                    {"role": "user", "content": text},
                ],
            },
            timeout=120.0,
        )
        response.raise_for_status()
        payload = response.json()
        return payload["choices"][0]["message"]["content"] or None
    except Exception:  # network, auth, rate limit — fall back, don't break import
        logger.warning("OpenRouter docx parse failed; falling back", exc_info=True)
        return None


def _anthropic_complete(api_key: str, model: str, text: str) -> str | None:
    """Call the Anthropic API directly. Returns the reply text, or None on any
    failure (including the SDK not being installed)."""
    try:
        import anthropic
    except ImportError:
        logger.info("anthropic SDK not installed; using heuristic docx parser")
        return None

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model=model,
            max_tokens=8192,
            system=_AI_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": text}],
        )
    except Exception:  # network, auth, rate limit — fall back, don't break import
        logger.warning("AI docx parse failed; falling back to heuristics", exc_info=True)
        return None

    return "".join(
        getattr(block, "text", "") for block in message.content
        if getattr(block, "type", None) == "text"
    ) or None


def _ai_parse(lines: list[str]) -> dict[str, Any] | None:
    """Ask an LLM to classify and split the exercise. Prefers OpenRouter when
    OPENROUTER_API_KEY is set, then the Anthropic API when ANTHROPIC_API_KEY is
    set. Returns None (so the caller falls back to heuristics) if no key is
    configured, the call fails, or the reply can't be used."""
    from app.core.config import settings

    text = "\n".join(lines).strip()
    if not text:
        return None

    raw: str | None = None
    openrouter_key = (getattr(settings, "openrouter_api_key", "") or "").strip()
    if openrouter_key:
        raw = _openrouter_complete(
            openrouter_key,
            getattr(settings, "openrouter_model", "anthropic/claude-haiku-4.5"),
            text,
            max_tokens=getattr(settings, "openrouter_max_tokens", 4096),
        )
    if raw is None:
        anthropic_key = (getattr(settings, "anthropic_api_key", "") or "").strip()
        if anthropic_key:
            raw = _anthropic_complete(
                anthropic_key,
                getattr(settings, "docx_ai_model", "claude-haiku-4-5"),
                text,
            )
    if raw is None:
        return None

    parsed = _extract_json(raw)
    if parsed is None:
        logger.warning("AI docx parse returned unusable output; falling back")
        return None
    return _normalise_ai(parsed)


def parse_docx(data: bytes) -> dict[str, Any]:
    lines = _extract_lines(data)

    # 1) Preferred: let Claude classify + split (when configured).
    ai = _ai_parse(lines)
    if ai:
        return ai

    # 2) Fallback heuristics. Pick a strategy from the rubric: a "match … with …"
    # exercise is parsed into left/right columns; everything else into rows.
    pre_instruction = ""
    for line in lines:
        if re.fullmatch(r"[\s_\.\-–—…]+", line):
            continue
        if _looks_instruction(line):
            pre_instruction = re.sub(r"^\(?\d{1,3}\s*[\.\)\-:]?\s+", "", line).strip()
            break
    if "match" in pre_instruction.lower():
        return _parse_matching(lines, pre_instruction)
    return _parse_repeat(lines)


def _parse_repeat(data_lines: list[str]) -> dict[str, Any]:
    lines = data_lines
    instructions = ""
    questions: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    for line in lines:
        # 0) Skip a blank "answer" line (underscores / dashes only), e.g. the
        #    space where the student rewrites a corrected sentence.
        if re.fullmatch(r"[\s_\.\-–—…]+", line):
            continue
        # 1) Exercise header / rubric — only before any question starts, so a
        #    question containing a phrase like "listen to" stays a question.
        if _looks_instruction(line) and current is None and not questions:
            if not instructions:
                instructions = re.sub(r"^\(?\d{1,3}\s*[\.\)\-:]?\s+", "", line).strip()
            continue

        # 2) Numbered question.
        qn = QNUM_RE.match(line)
        if qn:
            if current:
                questions.append(current)
            body = qn.group(2).strip()
            prefix, inline = _split_inline_options(body)
            # Options only count if they trail the prompt (prefix present); a bare
            # inline split with no prefix on a question line is almost always a
            # false positive, so keep the whole body as the prompt.
            if inline and prefix:
                current = {"prompt": prefix, "options": inline}
            else:
                current = {"prompt": body, "options": []}
            continue

        # 3) A line of grouped inline options ("a x  b y  c z").
        prefix, inline = _split_inline_options(line)
        if current is not None and not prefix and len(inline) >= 1:
            current["options"].extend(inline)
            continue

        # 4) A single option per line.
        delim = SINGLE_OPT_DELIM_RE.match(line)
        if current is not None and delim:
            current["options"].append(delim.group(2).strip())
            continue
        lone = LONE_OPT_RE.match(line)
        if current is not None and current["prompt"] and lone:
            current["options"].append(lone.group(2).strip())
            continue

        # 5) Continuation of a multi-line prompt. (A stray heading like a section
        #    title is ignored — instructions only come from rubric-looking lines.)
        if current is not None and not current["options"]:
            current["prompt"] = f"{current['prompt']} {line}".strip()

    if current:
        questions.append(current)

    for question in questions:
        question["options"] = [opt for opt in question["options"] if opt]
        # Normalise gap markers of any length to a consistent "___".
        question["prompt"] = re.sub(r"_{2,}", "___", question["prompt"])
    questions = [q for q in questions if q["prompt"]]

    # "Odd one out" / "wrong word in each group": the slash-separated words are
    # the answer choices, so turn each group into a multiple-choice question.
    instr_low = instructions.lower()
    is_word_group = any(k in instr_low for k in (
        "each group", "odd one out", "doesn't belong", "does not belong",
        "the wrong word", "which word", "different word", "the odd one",
    ))
    if is_word_group:
        for question in questions:
            parts = [p.strip() for p in re.split(r"\s*/\s*", question["prompt"]) if p.strip()]
            if len(parts) >= 2:
                question["options"] = parts
                question["prompt"] = "Choose the odd one out"
        template_key = "multiple_choice"
    else:
        template_key = _guess_type(instructions, questions)

    # For true/false statements the answer (T / F / NG) is written in a trailing
    # blank — strip it so the prompt is the clean statement.
    if template_key in ("true_false", "true_false_not_given"):
        for question in questions:
            prompt = re.sub(r"\s*_{2,}\s*$", "", question["prompt"])
            prompt = re.sub(r"\s+(?:NG|N/G|Not\s*Given|True|False|[TF])\s*$", "", prompt, flags=re.I)
            question["prompt"] = prompt.strip()

    return {
        "template_key": template_key,
        "instructions": instructions,
        "questions": questions,
    }


# ---------------------------------------------------------------------------
# Whole-document AI import (one .docx = a full test with many exercises)
# ---------------------------------------------------------------------------

_SEGMENT_PROMPT = """You segment an English test/coursebook document into its individual exercises and return STRICT JSON.

The input is the document as numbered lines: "N | text".

How these documents are laid out:
- An exercise starts with its number and a rubric line, e.g.
  "3 Recording 2 Listen again and complete the sentences ..." or
  "7 Choose the wrong word in each group."
- The exercise's numbered questions (1, 2, 3, ...) follow until the next
  exercise's rubric.
- A line containing only a score like "/5" or "/10" marks the END of the
  exercise it follows — include it in that exercise's range.
- Standalone section headers ("Listening", "Grammar", "Vocabulary",
  "Function", "Reading", "Writing") are NOT exercises; use them as the
  "section" of the exercises that follow.
- A final total such as "/100" belongs to no exercise.

Output shape (line numbers refer to the given numbering, end is inclusive):
{"exercises": [{"start": 1, "end": 11, "section": "Listening",
                "title": "Recording 1 multiple choice"}]}

Cover every exercise in the document, in order. Return ONLY the JSON object."""

_EXERCISE_PROMPT = """You extract ONE exercise from an English test/coursebook and return STRICT JSON.

The input has two parts: the FULL DOCUMENT (context only) and, after
"EXERCISE TO EXTRACT", the single exercise you must extract. Use the full
document only to find the reading passage the exercise refers to and to
determine answers — output questions ONLY for the exercise to extract.

Given the raw text of a single exercise:
1. Choose the best template_key from this exact list:
   multiple_choice, multi_select, true_false, true_false_not_given, gap_fill,
   dropdown_gap_fill, short_answer, error_correction, inline_alternatives,
   writing, speaking_prompt, matching_pairs, gap_match
2. Split it into individual questions with their prompts and options.
3. Extract or determine the correct answers:
   - EXAMPLES: the FIRST item of an exercise is often a pre-solved example.
     Signs that an item is an example: the answer word is already written in
     its gap ("She's Turkish."), a choice is already given after the question
     ("Fiona is going to  a ."), the corrected sentence is printed below the
     wrong one, the odd word is repeated after its group, or a matching
     letter is already filled in. Mark such an item "is_example": true, put
     the shown answer in correct_answer, and restore the prompt to its
     unanswered form (put ___ back where the example answer was written).
   - Otherwise determine the answer from English language knowledge (grammar,
     vocabulary) or from the reading passage and set "uncertain": false.
     Reading-comprehension answers MUST be worked out from the passage — read
     it carefully; they are not uncertain.
   - ONLY if the answer depends on an audio recording you cannot hear, give
     your best guess and set "uncertain": true.

Rules:
- Strip leading numbering ("1", "2)", "a.") from prompts and options.
- Keep blanks/gaps in prompts as exactly three underscores: ___
- NEVER include page numbers, "Recording 1", track numbers or score lines
  like "/5" in prompts, options or instructions.
- multiple_choice / multi_select: options are the choices without their a/b/c
  letters; correct_answer must EXACTLY equal one of the options (a list of
  them for multi_select).
- true_false: correct_answer is "True" or "False".
  true_false_not_given adds "Not Given".
- gap_fill / short_answer: correct_answer is the missing word or phrase. If
  the prompt gives the first letter ("the r..."), keep the prompt as
  "the r___" and give the full word as correct_answer.
- error_correction: prompt is the incorrect sentence, correct_answer is the
  corrected sentence.
- inline_alternatives: the prompt contains the two choices as
  "word1 / word2" with enough surrounding sentence; correct_answer is the
  correct word exactly as written. If one numbered item contains SEVERAL
  "word1 / word2" choices, split it into one question per choice.
- "odd one out" / "choose the wrong word" groups: use multiple_choice where
  options are the words of the group and correct_answer is the odd/wrong word.
- writing / speaking_prompt: the prompt is the full task text; no answers.
  If the rubric states a word count ("Write 50-70 words"), set min_words /
  max_words.
- matching_pairs: fill "left" (numbered prompts), "right" (lettered options)
  and "pairs" where pairs[i] is the letter ("a", "b", ...) of the right
  option matching left[i] (empty string if unknown).
- gap_match: a TWO-part exercise — "Complete the questions with the words in
  the box. THEN match questions 1-N with replies a-N". Use it whenever an
  exercise has BOTH a word box for gaps AND lettered replies to match. Fill:
  "words" = the words in the box (in the printed order, including the
  crossed-out example word), "right" = the lettered replies, and each
  question gets {"prompt": "... ___ ...", "word": "the box word for the gap",
  "reply": "the matching reply letter", "is_example": true when the word is
  pre-filled in the document}.
- If the exercise has a reading passage (email, article, story), put its full
  text in "passage" (plain text, keep paragraph breaks as newlines) and do
  NOT repeat it in prompts or instructions. If the exercise refers to a
  passage found elsewhere in the document ("Read the email again"), copy that
  passage into "passage" too, so the exercise is self-contained.
- "instructions" is the exercise rubric, cleaned of leading numbers.

Output shape (omit nothing; use [] / null / "" when not applicable):
{"template_key": "...", "instructions": "...", "passage": "",
 "min_words": null, "max_words": null,
 "questions": [{"prompt": "...", "options": [], "correct_answer": null,
                "is_example": false, "uncertain": false,
                "word": null, "reply": null}],
 "left": [], "right": [], "pairs": [], "words": []}

Return ONLY the JSON object. No markdown, no commentary."""


def ai_available() -> bool:
    from app.core.config import settings
    return bool((settings.openrouter_api_key or "").strip()
                or (settings.anthropic_api_key or "").strip())


async def _complete_async(system: str, text: str, max_tokens: int | None = None, retries: int = 4) -> str | None:
    """Async LLM call: OpenRouter when configured, else Anthropic. Retries
    transient failures (rate limits). Returns None when all attempts fail.
    Pass a small max_tokens for cheap short outputs (e.g. grading)."""
    from app.core.config import settings

    openrouter_key = (settings.openrouter_api_key or "").strip()
    anthropic_key = (settings.anthropic_api_key or "").strip()
    if max_tokens is None:
        max_tokens = getattr(settings, "openrouter_max_tokens", 4096)

    for attempt in range(retries):
        if attempt:
            await asyncio.sleep(5 * attempt)
        try:
            if openrouter_key:
                import httpx
                async with httpx.AsyncClient(timeout=180.0) as client:
                    response = await client.post(
                        "https://openrouter.ai/api/v1/chat/completions",
                        headers={"Authorization": f"Bearer {openrouter_key}"},
                        json={
                            "model": getattr(settings, "openrouter_model", "anthropic/claude-haiku-4.5"),
                            "max_tokens": max_tokens,
                            "messages": [
                                {"role": "system", "content": system},
                                {"role": "user", "content": text},
                            ],
                        },
                    )
                    # 402: OpenRouter reserves credit for max_tokens up front;
                    # a small balance can't cover concurrent reservations.
                    # Halve the reservation and retry.
                    if response.status_code == 402 and max_tokens > 512:
                        max_tokens = max(max_tokens // 2, 512)
                        logger.info("OpenRouter 402 — retrying with max_tokens=%s", max_tokens)
                        continue
                    # 429: free-tier models advertise how long to back off.
                    if response.status_code == 429 and attempt < 3:
                        retry_after = min(float(response.headers.get("Retry-After", 20)), 60.0)
                        logger.info("OpenRouter 429 — waiting %ss", retry_after)
                        await asyncio.sleep(retry_after)
                        continue
                    response.raise_for_status()
                    reply = response.json()["choices"][0]["message"]["content"]
                    if reply:
                        return reply
            elif anthropic_key:
                import anthropic
                client = anthropic.AsyncAnthropic(api_key=anthropic_key)
                message = await client.messages.create(
                    model=getattr(settings, "docx_ai_model", "claude-haiku-4-5"),
                    max_tokens=max_tokens,
                    system=system,
                    messages=[{"role": "user", "content": text}],
                )
                reply = "".join(
                    getattr(block, "text", "") for block in message.content
                    if getattr(block, "type", None) == "text"
                )
                if reply:
                    return reply
            else:
                return None
        except Exception:
            logger.warning("AI call failed (attempt %s)", attempt + 1, exc_info=True)
    return None


def _normalise_exercise(data: dict[str, Any] | None) -> dict[str, Any] | None:
    """Coerce a stage-2 reply into a clean exercise dict, or None."""
    if not isinstance(data, dict):
        return None
    template_key = str(data.get("template_key") or "").strip()
    if template_key not in _AI_ALLOWED:
        return None
    out: dict[str, Any] = {
        "template_key": template_key,
        "instructions": _clean_str(data.get("instructions")),
        "passage": str(data.get("passage") or "").strip(),
        "min_words": data.get("min_words") if isinstance(data.get("min_words"), int) else None,
        "max_words": data.get("max_words") if isinstance(data.get("max_words"), int) else None,
    }

    if template_key in _MATCHING_KEYS:
        left = [_clean_str(x) for x in (data.get("left") or []) if _clean_str(x)]
        right = [_clean_str(x) for x in (data.get("right") or []) if _clean_str(x)]
        if len(left) < 2 or len(right) < 2:
            return None
        pairs = [str(p or "").strip().lower() for p in (data.get("pairs") or [])]
        pairs += [""] * (len(left) - len(pairs))
        out.update({"left": left, "right": right, "pairs": pairs[: len(left)], "questions": []})
        return out

    if template_key == "gap_match":
        words = [_clean_str(x) for x in (data.get("words") or []) if _clean_str(x)]
        right = [_clean_str(x) for x in (data.get("right") or []) if _clean_str(x)]
        rows: list[dict[str, Any]] = []
        for item in data.get("questions") or []:
            if not isinstance(item, dict):
                continue
            prompt = _clean_str(item.get("prompt"))
            if not prompt:
                continue
            rows.append({
                "prompt": prompt,
                "word": _clean_str(item.get("word")),
                "reply": str(item.get("reply") or "").strip().lower(),
                "is_example": bool(item.get("is_example")),
            })
        if len(words) < 2 or len(right) < 2 or len(rows) < 2:
            return None
        out.update({"words": words, "right": right, "questions": rows})
        return out

    questions: list[dict[str, Any]] = []
    for item in data.get("questions") or []:
        if not isinstance(item, dict):
            continue
        prompt = _clean_str(item.get("prompt"))
        if not prompt:
            continue
        raw_ca = item.get("correct_answer")
        if isinstance(raw_ca, list):
            correct: Any = [_clean_str(x) for x in raw_ca if _clean_str(x)]
        else:
            correct = _clean_str(raw_ca) if raw_ca is not None else ""
        questions.append({
            "prompt": prompt,
            "options": [_clean_str(o) for o in (item.get("options") or []) if _clean_str(o)],
            "correct_answer": correct,
            "is_example": bool(item.get("is_example")),
            "uncertain": bool(item.get("uncertain")),
        })
    if not questions:
        return None
    out["questions"] = questions
    return out


async def ai_import_document(data: bytes) -> list[dict[str, Any]]:
    """Segment a full test .docx into exercises and parse each with the AI.
    Raises ValueError with a human-readable message when it cannot proceed."""
    lines = _extract_lines(data)
    if not lines:
        raise ValueError("The document is empty.")

    numbered = "\n".join(f"{i} | {line}" for i, line in enumerate(lines))
    raw = await _complete_async(_SEGMENT_PROMPT, numbered)
    seg = _extract_json(raw) if raw else None
    segments: list[dict[str, Any]] = []
    for item in (seg or {}).get("exercises") or []:
        try:
            start, end = int(item["start"]), int(item["end"])
        except (KeyError, TypeError, ValueError):
            continue
        if 0 <= start <= end < len(lines):
            segments.append({
                "start": start, "end": end,
                "section": str(item.get("section") or "").strip(),
                "title": str(item.get("title") or "").strip(),
            })
    if not segments:
        raise ValueError("The AI could not find any exercises in this document.")

    semaphore = asyncio.Semaphore(2)
    full_text = "\n".join(lines)

    async def parse_segment(segment: dict[str, Any]) -> dict[str, Any] | None:
        async with semaphore:
            chunk = "\n".join(lines[segment["start"]: segment["end"] + 1])
            prompt = f"FULL DOCUMENT (context only):\n{full_text}\n\nEXERCISE TO EXTRACT:\n{chunk}"
            reply = await _complete_async(_EXERCISE_PROMPT, prompt)
        exercise = _normalise_exercise(_extract_json(reply) if reply else None)
        if exercise is None:
            return None
        exercise["section"] = segment["section"]
        exercise["title"] = segment["title"]
        return exercise

    results = await asyncio.gather(*(parse_segment(s) for s in segments))
    exercises = [e for e in results if e]
    if not exercises:
        raise ValueError("The AI could not extract any usable exercises.")
    return exercises


_ANSWER_KEY_PROMPT = """You match an answer key document to a test's questions and return STRICT JSON.

You are given:
1. TEST STRUCTURE — the test's title and its exercises and questions. Every
   question has a short code like "e2q5" (exercise 2, question 5), its type,
   its prompt, and its lettered options when it has any.
2. ANSWER KEY — the raw text of an answer key document.

The answer key may contain the keys for SEVERAL tests (e.g. "MID-COURSE
TEST 1", "MID-COURSE TEST 2", "END-OF-COURSE TEST 1"). First find the section
that corresponds to the given test's title/variant number and use ONLY that
section.

Task: for every question whose answer you can find in the key (or determine
with certainty from it), output its code and the answer. The key's exercise
numbering follows the same order as the exercises in the structure.

Answer format by question type:
- multiple_choice: the EXACT text of the correct option. If the key gives a
  letter ("b"), convert it to that option's text using the options shown.
- multi_select: a JSON list of exact option texts.
- true_false: "True" or "False". true_false_not_given: may also be "Not Given".
- gap_fill / short_answer: the missing word or phrase.
- error_correction: the full corrected sentence.
- matching: the letter ("a", "b", ...) of the matching right-column option.
- Skip writing / speaking questions (they are graded by a teacher).

If you cannot confidently match an answer to a question, omit that code —
never guess from the question alone.

Output: {"answers": [{"code": "e1q1", "answer": "..."}]}
Return ONLY the JSON object. No markdown, no commentary."""


async def ai_match_answers(structure_text: str, key_text: str) -> list[dict[str, Any]]:
    """Ask the AI to map an answer-key document onto the test's questions.
    Returns a list of {"code", "answer"} dicts (possibly empty)."""
    prompt = f"TEST STRUCTURE:\n{structure_text}\n\nANSWER KEY:\n{key_text}"
    reply = await _complete_async(_ANSWER_KEY_PROMPT, prompt)
    data = _extract_json(reply) if reply else None
    results: list[dict[str, Any]] = []
    for item in (data or {}).get("answers") or []:
        if not isinstance(item, dict):
            continue
        code = str(item.get("code") or "").strip().lower()
        answer = item.get("answer")
        if code and answer is not None:
            results.append({"code": code, "answer": answer})
    return results


def normalise_answer(kind: str, options: list[Any], answer: Any,
                     tfng: bool = False) -> tuple[Any | None, str | None]:
    """Coerce an AI-matched answer into the exact stored form for its question
    kind. Returns (value, None) or (None, reason) when it can't be applied."""
    option_texts = [str(o).strip() for o in options]

    def match_option(value: Any) -> str | None:
        text = str(value).strip()
        if text in option_texts:
            return text
        low = text.lower()
        for opt in option_texts:
            if opt.lower() == low:
                return opt
        # A bare letter ("b") → that option.
        if len(text) == 1 and text.isalpha():
            idx = ord(text.lower()) - 97
            if 0 <= idx < len(option_texts):
                return option_texts[idx]
        return None

    if kind == "binary":
        allowed = ("True", "False", "Not Given") if tfng else ("True", "False")
        low = str(answer).strip().lower()
        for a in allowed:
            if low == a.lower():
                return a, None
        mapped = {"t": "True", "f": "False", "ng": "Not Given", "n/g": "Not Given"}.get(low)
        if mapped in allowed:
            return mapped, None
        return None, f"'{answer}' is not a valid {'/'.join(allowed)} answer"
    if kind == "options_multi":
        values = answer if isinstance(answer, list) else [answer]
        matched = [m for m in (match_option(v) for v in values) if m]
        if matched:
            return matched, None
        return None, f"'{answer}' does not match any option"
    if kind == "options_single":
        matched = match_option(answer)
        if matched is not None:
            return matched, None
        return None, f"'{answer}' does not match any option"
    if kind == "matching":
        text = str(answer).strip().lower()
        letters = {str(o.get("value") if isinstance(o, dict) else o).strip().lower() for o in options}
        if text in letters:
            return text, None
        return None, f"'{answer}' is not one of the match letters"
    # Free-text kinds (gap_fill, short_answer, error_correction).
    text = str(answer).strip()
    if isinstance(answer, list):
        text = ", ".join(str(x).strip() for x in answer if str(x).strip())
    if text:
        return text, None
    return None, "empty answer"


def _passage_to_html(text: str) -> str | None:
    paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
    if not paragraphs:
        return None
    return "".join(f"<p>{html_escape.escape(p)}</p>" for p in paragraphs)


_ALT_PAIR_RE = re.compile(r"([\w'’-]+)\s*/\s*([\w'’-]+)", re.UNICODE)


def build_task_payloads(exercises: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[str]]:
    """Convert parsed exercises into TaskCreate-shaped payloads that satisfy
    validate_task_payload. Fills conservative placeholders where the AI could
    not settle an answer and reports every such spot as a warning."""
    from app.services.question_templates import QUESTION_TEMPLATES

    letter = lambda i: chr(97 + i)  # noqa: E731
    payloads: list[dict[str, Any]] = []
    warnings: list[str] = []

    for index, exercise in enumerate(exercises, start=1):
        key = exercise["template_key"]
        template = QUESTION_TEMPLATES.get(key)
        if not template:
            warnings.append(f"Exercise {index}: unknown question type '{key}' — skipped.")
            continue
        per = template.get("per", [])
        title_bits = [b for b in (exercise.get("section"), exercise.get("title")) if b]
        title = (" · ".join(title_bits) or f"Exercise {index}")[:160]
        uncertain_count = 0

        interaction: dict[str, Any] = {"kind": template["kind"]}
        questions: list[dict[str, Any]] = []

        if template["mode"] == "repeat":
            for qi, q in enumerate(exercise["questions"], start=1):
                prompt = q["prompt"]
                options = q["options"]
                correct = q["correct_answer"]
                item: dict[str, Any] = {
                    "order_index": qi, "prompt": prompt, "points": 1,
                    "is_example": bool(q["is_example"]),
                }
                if q["uncertain"]:
                    uncertain_count += 1

                if "options" in per:
                    if len(options) < 2 and not item["is_example"]:
                        warnings.append(f"“{title}”: question {qi} has fewer than 2 options — skipped.")
                        continue
                    item["options"] = options
                    if "correct_multi" in per:
                        chosen = [c for c in (correct if isinstance(correct, list) else [correct]) if c in options]
                        if not chosen and options:
                            chosen = [options[0]]
                            warnings.append(f"“{title}”: question {qi} — answer not determined, defaulted to option a.")
                        item["correct_answer"] = chosen
                    else:
                        if correct not in options:
                            if options:
                                warnings.append(f"“{title}”: question {qi} — answer not determined, defaulted to option a.")
                                correct = options[0]
                        item["correct_answer"] = correct
                elif "correct_binary" in per or "correct_tfng" in per:
                    allowed = ("True", "False", "Not Given") if "correct_tfng" in per else ("True", "False")
                    low = str(correct).strip().lower()
                    normalized = next((a for a in allowed if low == a.lower()), None)
                    if normalized is None:
                        normalized = "True" if low in ("t", "true") else "False" if low in ("f", "false") \
                            else "Not Given" if low in ("ng", "n/g") and "correct_tfng" in per else None
                    if normalized is None:
                        normalized = "True"
                        if not item["is_example"]:
                            warnings.append(f"“{title}”: question {qi} — answer not determined, defaulted to True.")
                    item["correct_answer"] = normalized
                elif "correct_alt" in per:
                    match = _ALT_PAIR_RE.search(prompt)
                    if not match:
                        warnings.append(f"“{title}”: question {qi} has no 'word1 / word2' choice — skipped.")
                        continue
                    pair = [match.group(1), match.group(2)]
                    item["options"] = pair
                    item["correct_answer"] = correct if correct in pair else pair[0]
                    if correct not in pair and not item["is_example"]:
                        warnings.append(f"“{title}”: question {qi} — answer not determined, defaulted to the first word.")
                elif "correct_text" in per:
                    text_answer = str(correct).strip() if not isinstance(correct, list) else ""
                    if not text_answer and not item["is_example"]:
                        text_answer = "TODO"
                        warnings.append(f"“{title}”: question {qi} — answer missing, saved as TODO. Fix it before publishing.")
                    item["correct_answer"] = text_answer
                    item["case_sensitive"] = False
                    item["normalize_spaces"] = True

                if template.get("manual"):
                    item["correct_answer"] = None
                questions.append(item)

            if template.get("ex_words"):
                interaction["min_words"] = exercise.get("min_words") or 50
                interaction["max_words"] = exercise.get("max_words") or 200
                interaction["manual_review"] = True
            if template.get("ex_prep"):
                interaction["prep_seconds"] = 60
                interaction["manual_review"] = True

        elif template["mode"] == "composite:gapmatch":
            words, right = exercise["words"], exercise["right"]
            valid_letters = {letter(i) for i in range(len(right))}
            interaction["words"] = words
            interaction["options"] = [{"value": letter(i), "label": label} for i, label in enumerate(right)]
            interaction["reuse_options"] = False
            order = 0
            for qi, row in enumerate(exercise["questions"], start=1):
                word = row["word"]
                if word not in words:
                    warnings.append(f"“{title}”: word for question {qi} not from the box — check it.")
                    word = word or (words[0] if words else "")
                reply = row["reply"]
                if reply not in valid_letters:
                    reply = letter(min(qi - 1, len(right) - 1))
                    warnings.append(f"“{title}”: reply for question {qi} not determined — set to '{reply}'.")
                order += 1
                questions.append({
                    "order_index": order, "prompt": row["prompt"],
                    "correct_answer": word, "points": 1,
                    "is_example": bool(row["is_example"]),
                })
                order += 1
                questions.append({
                    "order_index": order, "prompt": f"Reply · {row['prompt']}",
                    "correct_answer": reply, "points": 1, "is_example": False,
                })
        elif template["mode"] == "composite:matching":
            left, right, pairs = exercise["left"], exercise["right"], exercise["pairs"]
            valid_letters = {letter(i) for i in range(len(right))}
            interaction["options"] = [{"value": letter(i), "label": label} for i, label in enumerate(right)]
            interaction["reuse_options"] = False
            for qi, prompt in enumerate(left):
                chosen = pairs[qi] if qi < len(pairs) else ""
                if chosen not in valid_letters:
                    chosen = letter(min(qi, len(right) - 1))
                    warnings.append(f"“{title}”: match for prompt {qi + 1} not determined — set to '{chosen}'.")
                questions.append({"order_index": qi + 1, "prompt": prompt, "correct_answer": chosen, "points": 1})
        else:
            warnings.append(f"“{title}”: type '{key}' is not supported by the importer — skipped.")
            continue

        if not questions:
            warnings.append(f"“{title}”: no usable questions — skipped.")
            continue

        if uncertain_count:
            warnings.append(f"“{title}”: {uncertain_count} answer(s) are AI guesses — double-check them.")
        instructions_low = exercise["instructions"].lower()
        if "listen" in instructions_low or "recording" in instructions_low:
            warnings.append(f"“{title}”: attach the audio file manually.")

        payloads.append({
            "title": title,
            "type": template["task_type"],
            "template_key": key,
            "instructions": exercise["instructions"] or title,
            "interaction": interaction,
            "passage_html": _passage_to_html(exercise.get("passage") or ""),
            "media_asset_id": None,
            "audio_replay_limit": None,
            "questions": questions,
        })

    return payloads, warnings
